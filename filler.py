"""
Заполнение шаблона Word (template.docx).

Плейсхолдеры: {{ date }}, {{ company }}, {{ contact_info }}
Тело письма вставляется перед параграфом 'С уважением,'.

fill_template(template_path, data, output_path)
  data = {
    "date": "30.04.2026",
    "company": "ООО Ромашка",
    "contact_info": "Менеджер: Баяндин О.Н.\nТел.: ___________",
    "body": "Текст тела письма от AI...",
  }
"""

import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_BODY_MARKER = 'С уважением,'

# Красная строка: 1.25 см → 709 твипс (1 см = 567 твипс)
_FIRST_LINE_TWIPS = 709

# Межстрочный интервал тела письма: 1.5 (276 twips = 240 * 1.15, или используем lineRule=auto + line=360 для 1.5)
_LINE_SPACING = 360  # 240 = single, 360 = 1.5, 480 = double


def _replace_in_paragraph(para, placeholder: str, replacement: str) -> bool:
    """Заменяет плейсхолдер в параграфе, учитывая разбивку по runs."""
    if placeholder not in para.text:
        return False

    full_text = para.text.replace(placeholder, replacement)

    for i, run in enumerate(para.runs):
        run.text = full_text if i == 0 else ''
    if not para.runs:
        para.add_run(full_text)
    return True


def _replace_in_doc(doc: Document, placeholder: str, replacement: str):
    """Заменяет плейсхолдер во всех параграфах и ячейках таблиц."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, placeholder, replacement)


def _make_body_paragraph(text: str) -> OxmlElement:
    """Создаёт параграф тела письма с красной строкой и межстрочным интервалом 1.5."""
    p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')

    # Красная строка (first-line indent)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), str(_FIRST_LINE_TWIPS))
    pPr.append(ind)

    # Межстрочный интервал 1.5
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(_LINE_SPACING))
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), '0')   # без дополнительного отступа после
    pPr.append(spacing)

    # Выравнивание по ширине (justified) — стандарт для деловых писем
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pPr.append(jc)

    p.append(pPr)

    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        if text.startswith(' ') or text.endswith(' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)

    return p


def _replace_body(doc: Document, body_text: str):
    """Вставляет тело письма перед 'С уважением,'.

    Разбивка: двойной перенос строки (\n\n) = новый абзац.
    Одиночный перенос внутри абзаца объединяется в один абзац.
    Каждый абзац оформляется с красной строкой и выравниванием по ширине.
    """
    target_para = None
    for para in doc.paragraphs:
        if _BODY_MARKER in para.text:
            target_para = para
            break

    if target_para is None:
        return

    # Разбиваем по двойному переносу → абзацы
    raw_blocks = body_text.split('\n\n')
    paragraphs = []
    for block in raw_blocks:
        # Одиночные переносы внутри блока объединяем через пробел
        clean = ' '.join(line.strip() for line in block.split('\n') if line.strip())
        if clean:
            paragraphs.append(clean)

    if not paragraphs:
        return

    # Вставляем абзацы перед 'С уважением,' в прямом порядке.
    # addprevious всегда вставляет перед target → каждый следующий
    # оказывается непосредственно перед target, итоговый порядок верный.
    for para_text in paragraphs:
        new_p = _make_body_paragraph(para_text)
        target_para._p.addprevious(new_p)

    # Пустая строка-разделитель перед 'С уважением,'
    spacer = _make_body_paragraph('')
    target_para._p.addprevious(spacer)


def fill_template(template_path: str, data: dict, output_path: str):
    """Заполняет шаблон данными и сохраняет результат."""
    doc = Document(template_path)

    _replace_in_doc(doc, '{{ date }}', data.get('date', ''))
    _replace_in_doc(doc, '{{ company }}', data.get('company', ''))
    _replace_in_doc(doc, '{{ contact_info }}', data.get('contact_info', ''))

    _replace_body(doc, data.get('body', ''))

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
